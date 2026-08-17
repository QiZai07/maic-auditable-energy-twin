from __future__ import annotations

import csv
import hashlib
import io
import math
import re
import shutil
import subprocess
import tempfile
import zipfile
from pathlib import Path
from typing import Any, Iterable

import pandas as pd


MAX_FILE_BYTES = 25 * 1024 * 1024
MAX_ARCHIVE_MEMBERS = 5_000
MAX_ARCHIVE_EXPANDED_BYTES = 200 * 1024 * 1024
MAX_PREVIEW_ROWS = 25

SUPPORTED_EXTENSIONS = {
    ".csv": ("Tabular data", "Phase 1"),
    ".xlsx": ("Excel workbook", "Phase 1"),
    ".xlsm": ("Excel workbook", "Phase 1"),
    ".pdf": ("PDF document", "Phase 2"),
    ".docx": ("Word document", "Phase 2"),
    ".png": ("Image or scan", "Phase 2"),
    ".jpg": ("Image or scan", "Phase 2"),
    ".jpeg": ("Image or scan", "Phase 2"),
    ".tif": ("Image or scan", "Phase 2"),
    ".tiff": ("Image or scan", "Phase 2"),
    ".dxf": ("CAD drawing", "Phase 3"),
    ".ifc": ("BIM model", "Phase 3"),
    ".dwg": ("CAD drawing", "Phase 3"),
}


FIELD_LIBRARY: dict[str, dict[str, Any]] = {
    "timestamp": {"label": "Timestamp", "unit": "ISO 8601", "aliases": ("timestamp", "date time", "datetime", "time", "recorded at", "reading date")},
    "billing_period": {"label": "Billing period", "unit": "YYYY-MM", "aliases": ("billing period", "bill month", "month", "period", "billing month")},
    "meter_id": {"label": "Meter identifier", "unit": "text", "aliases": ("meter id", "meter no", "meter number", "mpan", "account number", "point id")},
    "electricity_kwh": {"label": "Electricity consumption", "unit": "kWh", "aliases": ("electricity kwh", "energy kwh", "consumption kwh", "usage kwh", "electricity consumption", "active energy", "kwh")},
    "cumulative_kwh": {"label": "Cumulative meter reading", "unit": "kWh", "aliases": ("cumulative kwh", "meter reading", "total kwh", "register value", "cumulative energy")},
    "demand_kw": {"label": "Electrical demand", "unit": "kW", "aliases": ("demand kw", "load kw", "power kw", "peak demand", "maximum demand", "active power")},
    "grid_import_kwh": {"label": "Grid import", "unit": "kWh", "aliases": ("grid import", "import kwh", "purchased energy", "utility import")},
    "grid_export_kwh": {"label": "Grid export", "unit": "kWh", "aliases": ("grid export", "export kwh", "exported energy")},
    "pv_generation_kwh": {"label": "PV generation", "unit": "kWh", "aliases": ("pv generation", "solar generation", "pv energy", "photovoltaic generation", "solar kwh")},
    "tariff": {"label": "Electricity tariff", "unit": "currency/kWh", "aliases": ("tariff", "unit rate", "price per kwh", "energy rate", "electricity price")},
    "cost": {"label": "Electricity cost", "unit": "currency", "aliases": ("electricity cost", "energy cost", "bill amount", "cost", "charge", "total due")},
    "temperature_c": {"label": "Temperature", "unit": "degC", "aliases": ("temperature", "temperature c", "temp c", "outdoor temperature", "dry bulb")},
    "humidity_pct": {"label": "Relative humidity", "unit": "%", "aliases": ("humidity", "relative humidity", "rh", "humidity pct")},
    "ghi_w_m2": {"label": "Global horizontal irradiance", "unit": "W/m2", "aliases": ("ghi", "solar irradiance", "global horizontal irradiance", "irradiance")},
    "building_area_m2": {"label": "Gross floor area", "unit": "m2", "aliases": ("gross floor area", "floor area", "building area", "area m2", "gfa")},
    "equipment_id": {"label": "Equipment identifier", "unit": "text", "aliases": ("equipment id", "asset id", "plant id", "device id", "tag id")},
    "equipment_type": {"label": "Equipment type", "unit": "text", "aliases": ("equipment type", "asset type", "plant type", "device type", "category")},
    "capacity_kw": {"label": "Equipment capacity", "unit": "kW", "aliases": ("capacity kw", "rated power", "nameplate kw", "nominal capacity", "capacity")},
    "floor": {"label": "Floor or storey", "unit": "text", "aliases": ("floor", "storey", "story", "level")},
    "space": {"label": "Space or room", "unit": "text", "aliases": ("space", "room", "zone", "location", "area name")},
}


def _normalise_name(value: str) -> str:
    value = str(value).strip().lower()
    value = value.replace("²", "2").replace("°", "deg").replace("_", " ").replace("-", " ")
    return re.sub(r"[^a-z0-9%]+", " ", value).strip()


def _safe_value(value: Any) -> Any:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return None
    if isinstance(value, pd.Timestamp):
        return value.isoformat()
    if hasattr(value, "item"):
        return value.item()
    return str(value) if not isinstance(value, (str, int, float, bool)) else value


def _filename(name: str) -> str:
    clean = Path(str(name).replace("\\", "/")).name.strip()
    return re.sub(r"[^A-Za-z0-9._() -]", "_", clean)[:160] or "upload"


def _archive_guard(content: bytes) -> None:
    try:
        with zipfile.ZipFile(io.BytesIO(content)) as archive:
            members = archive.infolist()
            expanded = sum(member.file_size for member in members)
            if len(members) > MAX_ARCHIVE_MEMBERS or expanded > MAX_ARCHIVE_EXPANDED_BYTES:
                raise ValueError("The Office file expands beyond the safe processing limit.")
            for member in members:
                path = Path(member.filename)
                if path.is_absolute() or ".." in path.parts:
                    raise ValueError("The Office file contains an unsafe archive path.")
    except zipfile.BadZipFile as exc:
        raise ValueError("The Office file is not a valid ZIP-based document.") from exc


def validate_upload(name: str, content: bytes) -> tuple[str, str, str]:
    clean_name = _filename(name)
    extension = Path(clean_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        allowed = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type. Accepted extensions: {allowed}.")
    if not content:
        raise ValueError("The uploaded file is empty.")
    if len(content) > MAX_FILE_BYTES:
        raise ValueError(f"The file exceeds the {MAX_FILE_BYTES // 1024 // 1024} MB session limit.")
    if extension in {".xlsx", ".xlsm", ".docx"}:
        _archive_guard(content)
    if extension == ".pdf" and not content.startswith(b"%PDF"):
        raise ValueError("The file extension is PDF, but its signature does not match a PDF document.")
    if extension == ".png" and not content.startswith(b"\x89PNG\r\n\x1a\n"):
        raise ValueError("The file extension is PNG, but its signature does not match a PNG image.")
    kind, phase = SUPPORTED_EXTENSIONS[extension]
    return clean_name, kind, phase


def infer_field(column: str, series: pd.Series | None = None) -> dict[str, Any]:
    normalised = _normalise_name(column)
    tokens = set(normalised.split())
    best_id = "unmapped"
    best_score = 0.0
    for field_id, definition in FIELD_LIBRARY.items():
        for alias in definition["aliases"]:
            alias_normalised = _normalise_name(alias)
            alias_tokens = set(alias_normalised.split())
            score = 1.0 if normalised == alias_normalised else 0.0
            if alias_normalised and alias_normalised in normalised:
                score = max(score, 0.88)
            if alias_tokens:
                score = max(score, len(tokens & alias_tokens) / len(alias_tokens) * 0.78)
            if score > best_score:
                best_id, best_score = field_id, score

    if series is not None and len(series.dropna()):
        sample = series.dropna().head(100)
        if best_score < 0.72:
            parsed_dates = pd.to_datetime(sample, errors="coerce", utc=False)
            if parsed_dates.notna().mean() >= 0.9 and sample.astype(str).str.contains(r"[-/:T]", regex=True).mean() >= 0.5:
                best_id, best_score = "timestamp", 0.72

    if best_score < 0.48:
        return {"source": str(column), "target": "unmapped", "label": "Not mapped", "unit": "", "confidence": round(best_score, 2), "requires_confirmation": True}
    definition = FIELD_LIBRARY[best_id]
    return {
        "source": str(column),
        "target": best_id,
        "label": definition["label"],
        "unit": definition["unit"],
        "confidence": round(best_score, 2),
        "requires_confirmation": best_score < 0.85,
    }


def _unit_notes(column: str, mapping: dict[str, Any]) -> list[str]:
    name = _normalise_name(column)
    notes: list[str] = []
    target = mapping["target"]
    if target in {"electricity_kwh", "cumulative_kwh", "grid_import_kwh", "grid_export_kwh", "pv_generation_kwh"}:
        if re.search(r"(^| )mwh($| )", name):
            notes.append("Source appears to be MWh; multiply by 1,000 only after confirmation.")
        elif re.search(r"(^| )wh($| )", name) and "kwh" not in name:
            notes.append("Source appears to be Wh; divide by 1,000 only after confirmation.")
    if target in {"demand_kw", "capacity_kw"} and re.search(r"(^| )w($| )", name) and "kw" not in name:
        notes.append("Source appears to be W; divide by 1,000 only after confirmation.")
    if target == "temperature_c" and ("fahrenheit" in name or re.search(r"(^| )f($| )", name)):
        notes.append("Source appears to be degF; convert to degC only after confirmation.")
    return notes


def assess_table_quality(frame: pd.DataFrame, mappings: list[dict[str, Any]]) -> dict[str, Any]:
    issues: list[dict[str, str]] = []
    row_count = len(frame)
    duplicate_rows = int(frame.duplicated().sum()) if row_count else 0
    if duplicate_rows:
        issues.append({"severity": "warning", "check": "Duplicate rows", "detail": f"{duplicate_rows:,} duplicate row(s) detected."})

    for column in frame.columns:
        missing = int(frame[column].isna().sum())
        if row_count and missing:
            ratio = missing / row_count
            severity = "error" if ratio >= 0.5 else "warning"
            issues.append({"severity": severity, "check": "Missing values", "detail": f"{column}: {missing:,} missing ({ratio:.1%})."})

    mapped = {item["target"]: item["source"] for item in mappings if item["target"] != "unmapped"}
    for field in ("electricity_kwh", "demand_kw", "grid_import_kwh", "pv_generation_kwh", "cost"):
        column = mapped.get(field)
        if column is None:
            continue
        numeric = pd.to_numeric(frame[column], errors="coerce")
        negative = int((numeric < 0).sum())
        invalid = int(numeric.isna().sum() - frame[column].isna().sum())
        if invalid > 0:
            issues.append({"severity": "warning", "check": "Numeric format", "detail": f"{column}: {invalid:,} non-numeric value(s) require review."})
        if negative > 0:
            issues.append({"severity": "error", "check": "Negative value", "detail": f"{column}: {negative:,} negative reading(s) require confirmation."})
        clean = numeric.dropna()
        if len(clean) >= 8:
            q1, q3 = clean.quantile([0.25, 0.75])
            iqr = q3 - q1
            if iqr > 0:
                outliers = int(((clean < q1 - 3 * iqr) | (clean > q3 + 3 * iqr)).sum())
                if outliers:
                    issues.append({"severity": "warning", "check": "Extreme value", "detail": f"{column}: {outliers:,} value(s) fall beyond 3 IQR."})

    time_field = mapped.get("timestamp") or mapped.get("billing_period")
    coverage: dict[str, Any] = {"field": None, "start": None, "end": None, "granularity": "unknown", "valid_ratio": None}
    if time_field:
        dates = pd.to_datetime(frame[time_field], errors="coerce").dropna().sort_values()
        coverage["field"] = time_field
        coverage["valid_ratio"] = round(len(dates) / max(row_count, 1), 3)
        if len(dates):
            coverage["start"] = dates.iloc[0].isoformat()
            coverage["end"] = dates.iloc[-1].isoformat()
        if len(dates) >= 2:
            deltas = dates.diff().dropna().dt.total_seconds() / 60
            median_minutes = float(deltas.median())
            if median_minutes <= 20:
                coverage["granularity"] = "interval (15 minutes or finer)"
            elif median_minutes <= 90:
                coverage["granularity"] = "hourly"
            elif median_minutes <= 1_800:
                coverage["granularity"] = "daily"
            elif median_minutes <= 50_000:
                coverage["granularity"] = "monthly"
            duplicate_times = int(dates.duplicated().sum())
            if duplicate_times:
                issues.append({"severity": "warning", "check": "Duplicate timestamps", "detail": f"{duplicate_times:,} duplicate timestamp(s) detected."})

    errors = sum(item["severity"] == "error" for item in issues)
    warnings = sum(item["severity"] == "warning" for item in issues)
    score = max(0, 100 - errors * 18 - warnings * 5)
    return {"score": score, "errors": errors, "warnings": warnings, "issues": issues, "coverage": coverage}


def assess_readiness(tables: Iterable[dict[str, Any]], extracted_facts: Iterable[dict[str, Any]] = ()) -> dict[str, Any]:
    table_list = list(tables)
    targets = {
        mapping["target"]
        for table in table_list
        for mapping in table.get("mappings", [])
        if mapping.get("target") != "unmapped"
    }
    fact_types = {fact.get("type") for fact in extracted_facts}
    monthly = any(
        table.get("row_count", 0) >= 6
        and bool({"billing_period", "timestamp"} & {item.get("target") for item in table.get("mappings", [])})
        and bool({"electricity_kwh", "cumulative_kwh", "grid_import_kwh"} & {item.get("target") for item in table.get("mappings", [])})
        for table in table_list
    )
    interval = any(
        table.get("row_count", 0) >= 24
        and table.get("quality", {}).get("coverage", {}).get("granularity", "").startswith(("interval", "hourly"))
        for table in table_list
    )
    calibration_history = any(
        table.get("row_count", 0) >= 168
        and table.get("quality", {}).get("coverage", {}).get("granularity", "").startswith(("interval", "hourly"))
        for table in table_list
    )
    pv = "pv_generation_kwh" in targets
    weather = bool({"temperature_c", "humidity_pct", "ghi_w_m2"} & targets)
    assets = bool({"equipment_id", "equipment_type", "capacity_kw"} & targets or "equipment" in fact_types)
    building = bool("building_area_m2" in targets or "floor_area" in fact_types)

    capabilities = [
        {"name": "Monthly baseline and EUI", "ready": monthly and building, "needs": "A dated energy series and confirmed gross floor area"},
        {"name": "Tariff and bill screening", "ready": monthly and bool({"tariff", "cost"} & targets or "tariff" in fact_types), "needs": "Energy series plus tariff or bill amount"},
        {"name": "Interval anomaly detection", "ready": interval, "needs": "Hourly or finer timestamped energy or demand"},
        {"name": "PV performance assessment", "ready": pv and (weather or "ghi_w_m2" in targets), "needs": "PV generation plus irradiance or a confirmed benchmark"},
        {"name": "Operational calibration", "ready": calibration_history and weather, "needs": "At least one representative week of aligned hourly-or-finer load and weather"},
        {"name": "Equipment and space digital twin", "ready": assets and building, "needs": "Equipment register plus building or space data"},
    ]
    ready_count = sum(item["ready"] for item in capabilities)
    return {
        "score": round(ready_count / len(capabilities) * 100),
        "ready_count": ready_count,
        "total": len(capabilities),
        "capabilities": capabilities,
        "recommended_path": "Calibrate with interval data" if interval else "Establish the monthly baseline first" if monthly else "Complete field mapping and supply a dated energy series",
    }


def _table_record(name: str, frame: pd.DataFrame) -> dict[str, Any]:
    frame = frame.copy()
    frame.columns = [str(column).strip() or f"column_{index + 1}" for index, column in enumerate(frame.columns)]
    mappings = [infer_field(column, frame[column]) for column in frame.columns]
    unit_notes = [note for mapping in mappings for note in _unit_notes(mapping["source"], mapping)]
    preview = [{str(key): _safe_value(value) for key, value in row.items()} for row in frame.head(MAX_PREVIEW_ROWS).to_dict(orient="records")]
    return {
        "name": name,
        "row_count": len(frame),
        "column_count": len(frame.columns),
        "columns": list(frame.columns),
        "preview": preview,
        "mappings": mappings,
        "unit_notes": unit_notes,
        "quality": assess_table_quality(frame, mappings),
        # Kept only in session memory so reviewed mappings can be applied to the
        # complete table. Export helpers deliberately omit this private value.
        "_frame": frame,
    }


def _decode_text(content: bytes) -> tuple[str, str]:
    for encoding in ("utf-8-sig", "utf-8", "gb18030", "cp1252"):
        try:
            return content.decode(encoding), encoding
        except UnicodeDecodeError:
            pass
    raise ValueError("The text encoding could not be identified safely.")


def _parse_csv(content: bytes) -> tuple[list[dict[str, Any]], list[str]]:
    text, encoding = _decode_text(content)
    sample = text[:32_768]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        separator = dialect.delimiter
    except csv.Error:
        separator = ","
    frame = pd.read_csv(io.StringIO(text), sep=separator)
    return [_table_record("data", frame)], [f"Decoded as {encoding}; delimiter {separator!r}."]


def _parse_excel(content: bytes, extension: str) -> tuple[list[dict[str, Any]], list[str]]:
    workbook = pd.ExcelFile(io.BytesIO(content), engine="openpyxl")
    tables = [_table_record(sheet, workbook.parse(sheet_name=sheet)) for sheet in workbook.sheet_names]
    notes = [f"Workbook contains {len(tables)} sheet(s)."]
    if extension == ".xlsm":
        notes.append("Macro-enabled content was not executed; only worksheet values were read.")
    return tables, notes


def extract_text_facts(text: str) -> list[dict[str, Any]]:
    collapsed = re.sub(r"\s+", " ", text)
    facts: list[dict[str, Any]] = []
    patterns = [
        ("floor_area", r"(?:gross\s+floor\s+area|building\s+area|floor\s+area|gfa)\s*[:=]?\s*([\d,.]+)\s*(m2|m²|sqm|square\s+met(?:er|re)s?)", "m2"),
        ("energy", r"(?:electricity|energy|consumption|usage)\s*[:=]?\s*([\d,.]+)\s*(kwh|mwh)", None),
        ("demand", r"(?:peak\s+demand|maximum\s+demand|demand|load)\s*[:=]?\s*([\d,.]+)\s*(kw|mw)", None),
        ("capacity", r"(?:capacity|rated\s+power|installed\s+capacity)\s*[:=]?\s*([\d,.]+)\s*(kw|mw|kwp|mwP)", None),
        ("tariff", r"(?:tariff|unit\s+rate|price\s+per\s+kwh)\s*[:=]?\s*(?:[A-Z]{3}|[$¥£€])?\s*([\d,.]+)\s*(?:[A-Z]{3}|[$¥£€])?\s*(?:/|per)?\s*kwh", "currency/kWh"),
    ]
    for fact_type, pattern, fixed_unit in patterns:
        for match in re.finditer(pattern, collapsed, flags=re.IGNORECASE):
            value_text = match.group(1).replace(",", "")
            try:
                value = float(value_text)
            except ValueError:
                continue
            unit = fixed_unit or (match.group(2).lower() if match.lastindex and match.lastindex >= 2 else "")
            facts.append({"type": fact_type, "value": value, "unit": unit, "source_excerpt": collapsed[max(0, match.start() - 45): match.end() + 45].strip()})
    return facts[:100]


def _parse_pdf(content: bytes) -> tuple[str, dict[str, Any], list[str]]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ValueError("PDF support requires the pypdf package from requirements.txt.") from exc
    reader = PdfReader(io.BytesIO(content), strict=False)
    if reader.is_encrypted:
        raise ValueError("Password-protected PDFs must be unlocked before upload.")
    page_text = [(page.extract_text() or "").strip() for page in reader.pages]
    text = "\n".join(page_text)
    notes = [f"Extracted text from {sum(bool(item) for item in page_text)} of {len(page_text)} page(s)."]
    if len(text.strip()) < max(80, len(page_text) * 25):
        notes.append("Little embedded text was found. This may be a scanned PDF; optional cloud recognition can be used with explicit consent.")
    return text, {"pages": len(page_text), "text_characters": len(text)}, notes


def _parse_docx(content: bytes) -> tuple[str, dict[str, Any], list[str]]:
    try:
        from docx import Document
    except ImportError as exc:
        raise ValueError("Word support requires the python-docx package from requirements.txt.") from exc
    document = Document(io.BytesIO(content))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_rows = 0
    for table in document.tables:
        for row in table.rows:
            lines.append(" | ".join(cell.text.strip() for cell in row.cells))
            table_rows += 1
    text = "\n".join(lines)
    return text, {"paragraphs": len(document.paragraphs), "tables": len(document.tables), "table_rows": table_rows, "text_characters": len(text)}, ["Embedded macros and external links were not executed."]


def _parse_image(content: bytes) -> tuple[str, dict[str, Any], list[str]]:
    try:
        from PIL import Image
    except ImportError as exc:
        raise ValueError("Image support requires the Pillow package from requirements.txt.") from exc
    with Image.open(io.BytesIO(content)) as image:
        detail = {"width": image.width, "height": image.height, "format": image.format or "unknown", "mode": image.mode}
    return "", detail, ["Image validated locally. Text recognition is available only through the explicit cloud-enhancement control."]


def _parse_dxf(content: bytes) -> tuple[str, dict[str, Any], list[str]]:
    try:
        import ezdxf
    except ImportError as exc:
        raise ValueError("DXF support requires the ezdxf package from requirements.txt.") from exc
    with tempfile.TemporaryDirectory(prefix="irene-dxf-") as temp_dir:
        path = Path(temp_dir) / "drawing.dxf"
        path.write_bytes(content)
        document = ezdxf.readfile(path)
    modelspace = document.modelspace()
    entity_counts: dict[str, int] = {}
    layers: set[str] = set()
    text_items: list[str] = []
    blocks: list[str] = []
    closed_areas: list[float] = []
    for entity in modelspace:
        entity_type = entity.dxftype()
        entity_counts[entity_type] = entity_counts.get(entity_type, 0) + 1
        if hasattr(entity.dxf, "layer"):
            layers.add(str(entity.dxf.layer))
        if entity_type in {"TEXT", "MTEXT"}:
            value = entity.dxf.text if entity_type == "TEXT" else entity.plain_text()
            if value.strip():
                text_items.append(value.strip())
        elif entity_type == "INSERT":
            blocks.append(str(entity.dxf.name))
        elif entity_type == "LWPOLYLINE" and bool(entity.closed):
            try:
                from ezdxf.math import area

                closed_areas.append(abs(float(area(entity.get_points("xy")))))
            except Exception:
                pass
    header_units = int(document.header.get("$INSUNITS", 0) or 0)
    text = "\n".join(text_items)
    detail = {
        "dxf_version": document.dxfversion,
        "drawing_units_code": header_units,
        "layers": sorted(layers)[:250],
        "entity_counts": dict(sorted(entity_counts.items())),
        "block_references": sorted(set(blocks))[:250],
        "text_labels": text_items[:250],
        "closed_polyline_area_values": closed_areas[:250],
    }
    notes = ["Drawing entities, layers, blocks and labels were extracted locally."]
    if header_units == 0:
        notes.append("Drawing units are unspecified; derived dimensions or areas require manual unit confirmation.")
    return text, detail, notes


def _parse_ifc(content: bytes) -> tuple[str, dict[str, Any], list[str]]:
    text, encoding = _decode_text(content)
    if "ISO-10303-21" not in text[:1_000].upper() or "IFC" not in text[:5_000].upper():
        raise ValueError("The IFC header could not be validated.")
    entity_names = re.findall(r"=\s*(IFC[A-Z0-9_]+)\s*\(", text, flags=re.IGNORECASE)
    counts: dict[str, int] = {}
    for name in entity_names:
        key = name.upper()
        counts[key] = counts.get(key, 0) + 1
    labels: list[str] = []
    for entity in ("IFCBUILDING", "IFCBUILDINGSTOREY", "IFCSPACE", "IFCSYSTEM", "IFCEQUIPMENTELEMENT", "IFCFLOWTERMINAL", "IFCENERGYCONVERSIONDEVICE"):
        for match in re.finditer(rf"{entity}\s*\(([^;]{{0,1200}})\);", text, flags=re.IGNORECASE):
            quoted = re.findall(r"'((?:''|[^'])*)'", match.group(1))
            labels.extend(item.replace("''", "'") for item in quoted[1:3] if item and item != "$")
    detail = {
        "encoding": encoding,
        "schema": next(iter(re.findall(r"FILE_SCHEMA\s*\(\s*\(\s*'([^']+)'", text, flags=re.IGNORECASE)), "unknown"),
        "entity_count": len(entity_names),
        "entity_counts": dict(sorted(counts.items(), key=lambda item: (-item[1], item[0]))[:200]),
        "buildings": counts.get("IFCBUILDING", 0),
        "storeys": counts.get("IFCBUILDINGSTOREY", 0),
        "spaces": counts.get("IFCSPACE", 0),
        "equipment_entities": sum(count for name, count in counts.items() if any(token in name for token in ("EQUIPMENT", "TERMINAL", "BOILER", "CHILLER", "PUMP", "FAN"))),
        "labels": labels[:250],
    }
    return "\n".join(labels), detail, ["IFC STEP entities and object labels were extracted locally; geometry is not recalculated in this lightweight review."]


def dwg_support_status() -> dict[str, Any]:
    configured = shutil.which("ODAFileConverter") or shutil.which("ODAFileConverter.exe") or shutil.which("dwg2dxf")
    return {
        "available": bool(configured),
        "converter": configured,
        "message": "A local DWG-to-DXF converter is available." if configured else "No local DWG converter was found. Export the drawing as DXF or configure ODA File Converter / LibreDWG locally.",
    }


def _parse_dwg(content: bytes) -> tuple[str, dict[str, Any], list[str]]:
    signature = content[:6].decode("ascii", errors="replace")
    if not re.fullmatch(r"AC10\d{2}", signature):
        raise ValueError("The DWG signature could not be validated.")
    support = dwg_support_status()
    detail = {"dwg_signature": signature, "conversion": support}
    notes = [support["message"], "The original DWG was not modified or persisted."]
    return "", detail, notes


def parse_client_file(name: str, content: bytes) -> dict[str, Any]:
    clean_name, kind, phase = validate_upload(name, content)
    extension = Path(clean_name).suffix.lower()
    tables: list[dict[str, Any]] = []
    notes: list[str] = []
    extracted_text = ""
    details: dict[str, Any] = {}
    if extension == ".csv":
        tables, notes = _parse_csv(content)
    elif extension in {".xlsx", ".xlsm"}:
        tables, notes = _parse_excel(content, extension)
    elif extension == ".pdf":
        extracted_text, details, notes = _parse_pdf(content)
    elif extension == ".docx":
        extracted_text, details, notes = _parse_docx(content)
    elif extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"}:
        extracted_text, details, notes = _parse_image(content)
    elif extension == ".dxf":
        extracted_text, details, notes = _parse_dxf(content)
    elif extension == ".ifc":
        extracted_text, details, notes = _parse_ifc(content)
    elif extension == ".dwg":
        extracted_text, details, notes = _parse_dwg(content)

    facts = extract_text_facts(extracted_text)
    readiness = assess_readiness(tables, facts)
    status = "review_required"
    if extension == ".dwg" and not details.get("conversion", {}).get("available"):
        status = "conversion_required"
    elif extension in {".png", ".jpg", ".jpeg", ".tif", ".tiff"} or (extension == ".pdf" and details.get("text_characters", 0) < 80):
        status = "recognition_optional"
    return {
        "filename": clean_name,
        "extension": extension,
        "kind": kind,
        "phase": phase,
        "size_bytes": len(content),
        "sha256": hashlib.sha256(content).hexdigest(),
        "status": status,
        "tables": tables,
        "extracted_facts": facts,
        "details": details,
        "notes": notes,
        "readiness": readiness,
        "confirmation_required": True,
        "privacy": "Processed in the current session. No raw file is retained by this parser.",
    }


def build_mapping_template(manifest: dict[str, Any]) -> dict[str, Any]:
    return {
        "schema_version": "irene-client-data-v1",
        "source": {
            "filename": manifest["filename"],
            "sha256": manifest["sha256"],
            "kind": manifest["kind"],
        },
        "tables": [
            {
                "name": table["name"],
                "mappings": [
                    {"source": mapping["source"], "target": mapping["target"], "unit": mapping["unit"], "confirmed": False}
                    for mapping in table["mappings"]
                ],
            }
            for table in manifest.get("tables", [])
        ],
        "extracted_facts": [dict(fact, confirmed=False) for fact in manifest.get("extracted_facts", [])],
        "approved_for_model": False,
    }


def convert_dwg_to_dxf(content: bytes) -> bytes:
    support = dwg_support_status()
    converter = support.get("converter")
    if not converter:
        raise RuntimeError(support["message"])
    with tempfile.TemporaryDirectory(prefix="irene-dwg-") as temp_dir:
        base = Path(temp_dir)
        source = base / "source.dwg"
        output = base / "output"
        output.mkdir()
        source.write_bytes(content)
        executable = Path(converter).name.lower()
        if "oda" in executable:
            command = [converter, str(base), str(output), "ACAD2018", "DXF", "0", "1", "source.dwg"]
        else:
            command = [converter, str(source), str(output / "source.dxf")]
        completed = subprocess.run(command, capture_output=True, text=True, timeout=90, check=False)
        candidates = list(output.glob("*.dxf")) + list(base.glob("*.dxf"))
        if completed.returncode != 0 or not candidates:
            message = (completed.stderr or completed.stdout or "converter returned no DXF output").strip()[-500:]
            raise RuntimeError(f"DWG conversion failed: {message}")
        return candidates[0].read_bytes()
